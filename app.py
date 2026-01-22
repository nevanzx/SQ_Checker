import streamlit as st
import os
import json
import tempfile
from docx import Document
from docx.shared import Inches
from datetime import datetime
import requests
import time
import concurrent.futures
from threading import Lock

# Import the prompts module
from prompts import (
    get_deepseek_prompt
)

# Function to load API keys from key.json
def load_api_keys():
    try:
        with open('key.json', 'r') as f:
            keys_data = json.load(f)

        # Create a mapping from service names to API keys
        key_map = {}
        for service in keys_data.get('apis', []):
            name = service['name']
            keys = service['keys']
            # Use the first key if multiple keys are provided
            key_map[name] = keys[0] if keys else ""

        return key_map
    except FileNotFoundError:
        # If key.json doesn't exist, return an empty mapping
        return {}
    except Exception as e:
        st.error(f"Error loading API keys: {e}")
        return {}

def main():
    # Load API keys from key.json
    api_keys = load_api_keys()

    # Initialize session state
    if 'models' not in st.session_state:
        st.session_state.models = [
            {"name": "DeepSeek Reasoner", "api_key": api_keys.get('deepseek', ''), "provider": "deepseek", "temperature": 0.3}
        ]

    if 'uploaded_files' not in st.session_state:
        st.session_state.uploaded_files = []

    if 'analysis_results' not in st.session_state:
        st.session_state.analysis_results = []

    st.set_page_config(
        page_title="Survey Questionnaire Quality Checker",
        page_icon="📋",
        layout="wide"
    )

    st.title("📋 Survey Questionnaire Quality Checker")
    st.markdown("""
    This application analyzes survey questionnaires using multiple AI models to assess quality metrics.
    Upload your survey files, select AI models, and get detailed quality reports.
    """)

    # Create tabs for different sections - removed Model Management tab
    tab1, tab2 = st.tabs(["Upload & Results", "Settings"])

    with tab1:
        upload_and_results_section()

    with tab2:
        settings_section()

def upload_and_results_section():
    # API key upload section (at the top)
    st.subheader("Upload API Keys File")
    uploaded_key_file = st.file_uploader(
        "Upload key.json file to update all API keys",
        type=['json'],
        key="key_file_uploader_combined"
    )

    if uploaded_key_file is not None:
        try:
            # Load the uploaded JSON file
            uploaded_keys = json.load(uploaded_key_file)

            # Validate the structure of the uploaded file
            if 'apis' not in uploaded_keys:
                st.error("Uploaded file is not a valid key.json format. Missing 'apis' key.")
            else:
                # Update the session state models with the uploaded keys
                for service in uploaded_keys['apis']:
                    service_name = service['name']
                    service_keys = service['keys']

                    # Update the corresponding model in session state
                    for model in st.session_state.models:
                        if service_name == 'deepseek' and model['provider'] == 'deepseek':
                            model['api_key'] = service_keys[0] if service_keys else ""

                # Save the uploaded keys to the local key.json file
                with open('key.json', 'w') as f:
                    json.dump(uploaded_keys, f, indent=2)

                st.success("API keys updated successfully from uploaded file!")
                # Removed st.rerun() to prevent the file uploader from flashing
                # The API keys are updated in session state and will be used in subsequent calls
        except json.JSONDecodeError:
            st.error("Uploaded file is not a valid JSON file.")
        except Exception as e:
            st.error(f"Error processing uploaded file: {e}")

    st.markdown("---")  # Divider

    # Upload section
    st.header("Upload Survey Questionnaires")

    uploaded_files = st.file_uploader(
        "Choose survey questionnaire files (TXT, JSON, CSV, DOCX, PDF)",
        type=['txt', 'json', 'csv', 'docx', 'pdf'],
        accept_multiple_files=True
    )

    if uploaded_files:
        st.session_state.uploaded_files = uploaded_files
        st.success(f"Uploaded {len(uploaded_files)} file(s)")

        # Display uploaded files
        for i, file in enumerate(st.session_state.uploaded_files):
            st.write(f"{i+1}. {file.name}")

    # Model selection
    st.header("Select AI Model for Analysis")

    # Create a list of model names for the dropdown
    model_names = [model['name'] for model in st.session_state.models]
    selected_model_name = st.selectbox("Choose an AI model for analysis", options=model_names, index=0)

    # Get the selected model object based on the name
    selected_model = next((model for model in st.session_state.models if model['name'] == selected_model_name), None)

    # Temperature control for the selected model
    if selected_model:
        st.subheader(f"Temperature for {selected_model['name']}")
        temperature = st.slider(
            f"Set temperature for {selected_model['name']}",
            min_value=0.0,
            max_value=1.0,
            value=selected_model['temperature'],
            step=0.05,
            key=f"temp_{selected_model['provider']}"
        )
        # Update the temperature in the session state
        selected_model['temperature'] = temperature
        st.caption("Temperature controls randomness. Lower values make responses more deterministic, higher values more creative.")

    # Analyze button
    if st.button("Analyze Survey Quality", type="primary"):
        if not st.session_state.uploaded_files:
            st.error("Please upload at least one survey questionnaire file")
        elif not selected_model:
            st.error("Please select an AI model")
        else:
            # Show loading status with detailed text
            # Show progress information
            progress_text = f"Starting analysis of {len(st.session_state.uploaded_files)} file(s) with {selected_model_name}..."
            st.info(progress_text)

            with st.spinner(f"Analyzing survey quality with {selected_model_name}..."):
                analyze_surveys([selected_model])

    # Results section
    st.header("Analysis Results")

    if st.session_state.analysis_results:
        st.success(f"Found {len(st.session_state.analysis_results)} analysis result(s)")

        for i, result in enumerate(st.session_state.analysis_results):
            with st.expander(f"Result {i+1}: {result['filename']}"):
                # Auto-generate DOCX file and provide download link
                docx_file = generate_docx(result['analysis'], result['filename'])
                with open(docx_file, "rb") as f:
                    st.download_button(
                        label="Download DOCX Report",
                        data=f,
                        file_name=f"quality_report_{result['filename']}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"docx_download_{i}"
                    )

                # Optional: Show a preview of the analysis (first few lines)
                with st.popover("View Analysis Summary"):
                    st.write("**Overall Assessment:**")
                    assessment = result['analysis'].get('overall_assessment', 'No assessment available')
                    # Limit the display to first 500 characters for preview
                    st.caption(assessment[:500] + "..." if len(assessment) > 500 else assessment)

    else:
        st.info("No analysis results yet. Upload surveys and run analysis to see results here.")



def settings_section():
    st.header("Application Settings")

    st.subheader("Export Options")
    include_charts = st.checkbox("Include charts in DOCX export", value=True)
    detailed_analysis = st.checkbox("Include detailed analysis in export", value=True)

    st.subheader("AI Model Settings")
    # Allow setting default temperature for all models
    default_temperature = st.slider(
        "Default temperature for all AI models",
        min_value=0.0,
        max_value=1.0,
        value=0.3,
        step=0.1,
        key="default_temp_slider"
    )

    if st.button("Apply to all models"):
        for model in st.session_state.models:
            model['temperature'] = default_temperature
        st.success(f"Default temperature {default_temperature} applied to all models!")

    st.subheader("Quality Metrics")
    metrics = st.multiselect(
        "Select quality metrics to analyze:",
        ["Clarity", "Bias Detection", "Relevance", "Completeness", "Logical Flow", "Response Options", "Demographic Appropriateness"],
        default=["Clarity", "Bias Detection", "Relevance"]
    )

def analyze_single_file(uploaded_file, selected_models):
    """Analyze a single survey file using selected AI models - returns analysis without UI updates"""
    # Process different file types
    file_content = process_uploaded_file(uploaded_file)

    if not file_content:
        return {'filename': uploaded_file.name, 'error': f"Could not process file: {uploaded_file.name}"}

    # Prepare analysis for each selected model
    file_analysis = {
        'filename': uploaded_file.name,
        'models_used': [],
        'individual_question_analysis': [],
        'recommendations': [],
        'overall_assessment': "",
        'timestamp': datetime.now().isoformat()
    }

    for model in selected_models:
        # Call AI model for analysis
        model_analysis = call_ai_model(file_content, model)
        file_analysis['models_used'].append({
            'model_name': model['name'],
            'analysis': model_analysis
        })

        # Process model analysis results
        if 'recommendations' in model_analysis:
            file_analysis['recommendations'].extend(model_analysis['recommendations'])

        # Include detailed analysis components if present
        if 'individual_question_analysis' in model_analysis:
            file_analysis['individual_question_analysis'].extend(model_analysis['individual_question_analysis'])

        if 'overall_assessment' in model_analysis:
            # Clean up the model assessment to remove any JSON formatting
            raw_assessment = model_analysis['overall_assessment']
            clean_model_assessment = raw_assessment
            if raw_assessment:
                # Check if the raw assessment looks like a complete JSON response
                # (starts with { and ends with }, which would indicate the entire response is JSON)
                stripped = raw_assessment.strip()
                if stripped.startswith('{') and stripped.endswith('}'):
                    # This looks like the entire response is JSON, which means the model returned
                    # the full JSON structure as the overall assessment
                    # Try to parse it and extract just the actual assessment text
                    try:
                        parsed = json.loads(raw_assessment)
                        # If it has an overall_assessment field, use that
                        if 'overall_assessment' in parsed and isinstance(parsed['overall_assessment'], str):
                            clean_model_assessment = parsed['overall_assessment']
                        else:
                            # If not, just clean up the JSON formatting markers
                            clean_model_assessment = raw_assessment.replace('```json', '').replace('```', '').strip()
                            import re
                            clean_model_assessment = re.sub(r'\s+', ' ', clean_model_assessment)
                    except json.JSONDecodeError:
                        # If it's not valid JSON, just clean up formatting markers
                        clean_model_assessment = raw_assessment.replace('```json', '').replace('```', '').strip()
                        import re
                        clean_model_assessment = re.sub(r'\s+', ' ', clean_model_assessment)
                else:
                    # Just clean up formatting markers for regular text
                    clean_model_assessment = raw_assessment.replace('```json', '').replace('```', '').strip()
                    import re
                    clean_model_assessment = re.sub(r'\s+', ' ', clean_model_assessment)

            if file_analysis['overall_assessment']:
                file_analysis['overall_assessment'] += f"\n\n{clean_model_assessment}"
            else:
                file_analysis['overall_assessment'] = f"{clean_model_assessment}"

    return {
        'filename': uploaded_file.name,
        'analysis': file_analysis
    }

def analyze_surveys(selected_models):
    """Analyze uploaded surveys using selected AI models with parallel processing"""
    results = []

    # Show overall progress
    total_files = len(st.session_state.uploaded_files)
    progress_text = f"Starting parallel analysis of {total_files} file(s)..."
    st.info(progress_text)

    # Use ThreadPoolExecutor for parallel processing
    with concurrent.futures.ThreadPoolExecutor(max_workers=min(4, len(st.session_state.uploaded_files))) as executor:
        # Submit all file processing tasks
        future_to_file = {
            executor.submit(analyze_single_file, uploaded_file, selected_models): uploaded_file
            for uploaded_file in st.session_state.uploaded_files
        }

        # Collect results as they complete
        completed_count = 0
        for future in concurrent.futures.as_completed(future_to_file):
            completed_count += 1
            uploaded_file = future_to_file[future]

            # Update progress
            progress_text = f"Completed {completed_count} of {total_files} files: {uploaded_file.name}"
            st.info(progress_text)

            try:
                result = future.result()
                if result and 'error' not in result:
                    results.append(result)
                elif 'error' in result:
                    st.error(result['error'])
            except Exception as e:
                st.error(f"Error processing file {uploaded_file.name}: {str(e)}")

    st.session_state.analysis_results = results
    st.success(f"Parallel analysis complete for {len(results)} file(s)!")

def process_uploaded_file(uploaded_file):
    """Process different file types and extract text content, including tables"""
    import fitz  # PyMuPDF for PDF processing
    import csv
    from io import StringIO, BytesIO

    file_extension = uploaded_file.name.split('.')[-1].lower()

    try:
        if file_extension == 'txt':
            return uploaded_file.getvalue().decode("utf-8")

        elif file_extension == 'json':
            content = uploaded_file.getvalue().decode("utf-8")
            json_data = json.loads(content)
            # Convert JSON to string for analysis
            return json.dumps(json_data, indent=2)

        elif file_extension == 'csv':
            string_data = StringIO(uploaded_file.getvalue().decode("utf-8"))
            csv_reader = csv.reader(string_data)
            content = '\n'.join([','.join(row) for row in csv_reader])
            return content

        elif file_extension == 'docx':
            from docx import Document
            doc = Document(BytesIO(uploaded_file.getvalue()))

            # Extract text content (images are automatically ignored by python-docx)
            content = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])

            # Extract table content
            table_content = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_content.append('| ' + ' | '.join(row_data) + ' |')

            # Combine text and table content
            if table_content:
                content += "\n\nExtracted Tables:\n" + '\n'.join(table_content)

            return content

        elif file_extension == 'pdf':
            pdf_document = fitz.open(stream=uploaded_file.getvalue(), filetype="pdf")
            content = ""

            for page_num in range(pdf_document.page_count):
                page = pdf_document.load_page(page_num)

                # Extract regular text
                content += page.get_text() + "\n"

                # Extract tables using PyMuPDF's table functionality
                try:
                    tables = page.find_tables()
                    for i, table in enumerate(tables):
                        content += f"\nTable {i+1}:\n"
                        for row in table.rows:
                            row_text = " | ".join([cell.text.strip() for cell in row.cells])
                            content += f"{row_text}\n"
                        content += "\n"
                except:
                    # If table extraction fails, continue with just text
                    pass

            pdf_document.close()
            return content

        else:
            print(f"Unsupported file type: {file_extension}")
            return None

    except Exception as e:
        print(f"Error processing file {uploaded_file.name}: {str(e)}")
        return None

def call_ai_model(file_content, model):
    """Call the DeepSeek AI model for analysis"""
    headers = {}
    payload = {}

    if model['provider'] == 'deepseek':
        # DeepSeek Reasoner API call
        headers = {
            'Authorization': f"Bearer {model['api_key']}",
            'Content-Type': 'application/json'
        }

        # Get the messages list from the updated prompt function
        messages = get_deepseek_prompt(file_content)
        payload = {
            "model": "deepseek-reasoner",
            "messages": messages,  # This is now a list of messages with system and user roles
            "response_format": {"type": "json_object"},
            "temperature": model.get('temperature', 0.3)
        }
        url = model.get('endpoint') or 'https://api.deepseek.com/chat/completions'

    try:
        response = requests.post(url, headers=headers, json=payload)
        response.raise_for_status()
        result = response.json()

        # Extract the content from the response
        content = result['choices'][0]['message']['content']

        # Try to extract valid JSON from the content
        analysis = extract_valid_json(content)

        # If we couldn't extract valid JSON, wrap it in our expected format
        if analysis is None:
            analysis = {
                "individual_question_analysis": [],
                "overall_assessment": content,
                "recommendations": ["This model did not return structured JSON. Raw analysis: " + content]
            }

        # Clean up any JSON formatting that might be embedded in the overall assessment
        if 'overall_assessment' in analysis and analysis['overall_assessment']:
            # Remove any JSON code block markers and clean up the text
            assessment = analysis['overall_assessment']
            # Remove markdown code block markers if present
            assessment = assessment.replace('```json', '').replace('```', '').strip()
            # If the assessment looks like it's just JSON, try to extract meaningful text
            if assessment.startswith('{') and assessment.endswith('}'):
                # This means the entire assessment field was returned as JSON, which shouldn't happen
                # The assessment should be plain text, not JSON structure
                extracted = extract_valid_json(assessment)
                if extracted and 'overall_assessment' in extracted:
                    analysis['overall_assessment'] = extracted['overall_assessment']

        return analysis
    except Exception as e:
        return {
            "individual_question_analysis": [],
            "overall_assessment": "",
            "recommendations": [f"Error analyzing with {model['name']}: {str(e)}"]
        }

def extract_valid_json(content):
    """
    Extract valid JSON from content that may contain additional text or formatting.
    This function tries multiple strategies to extract valid JSON from AI model responses.
    """
    import re
    import json

    # Strategy 1: Try to parse as-is
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        pass

    # Strategy 2: Remove markdown code blocks and try to parse
    cleaned_content = content.replace('```json', '').replace('```', '').strip()
    try:
        return json.loads(cleaned_content)
    except json.JSONDecodeError:
        pass

    # Strategy 3: Extract JSON object using regex (look for content between { and })
    # This handles cases where there's text before or after the JSON
    json_match = re.search(r'\{.*\}', content, re.DOTALL)
    if json_match:
        json_str = json_match.group()
        try:
            return json.loads(json_str)
        except json.JSONDecodeError:
            pass

    # Strategy 4: Look for JSON array as well (for cases where the response is an array)
    array_match = re.search(r'\[.*\]', content, re.DOTALL)
    if array_match:
        array_str = array_match.group()
        try:
            return json.loads(array_str)
        except json.JSONDecodeError:
            pass

    # If all strategies fail, return None
    return None


def generate_docx(analysis_data, filename):
    """Generate a DOCX report from analysis data"""
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # Set document title with larger font
    title = doc.add_heading('Survey Questionnaire Quality Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.bold = True

    # Add metadata section
    metadata_section = doc.add_paragraph()
    metadata_section.add_run(f'Analyzed File: ').bold = True
    metadata_section.add_run(f'{filename}')
    metadata_section.add_run('\nGenerated on: ').bold = True
    metadata_section.add_run(f'{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

    # Add a horizontal line
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.add_run().add_break()

    # Add executive summary section
    doc.add_heading('Executive Summary', level=1)

    # Count valid and invalid questions
    valid_count = 0
    invalid_count = 0
    total_questions = 0

    if 'individual_question_analysis' in analysis_data and analysis_data['individual_question_analysis']:
        for question in analysis_data['individual_question_analysis']:
            validity = question.get('validity', 'N/A')
            if validity.lower() == 'valid':
                valid_count += 1
            elif validity.lower() == 'not valid':
                invalid_count += 1
            total_questions += 1

    summary_para = doc.add_paragraph()
    summary_para.add_run(f'Total Questions Analyzed: ').bold = True
    summary_para.add_run(f'{total_questions}\n')
    summary_para.add_run(f'Valid Questions: ').bold = True
    summary_para.add_run(f'{valid_count}\n')
    summary_para.add_run(f'Invalid Questions: ').bold = True
    summary_para.add_run(f'{invalid_count}\n')

    if total_questions > 0:
        valid_percentage = (valid_count / total_questions) * 100
        summary_para.add_run(f'Quality Score: ').bold = True
        summary_para.add_run(f'{valid_percentage:.1f}%\n')
    else:
        valid_percentage = 0  # Default value when no questions are analyzed

    # Add individual question analysis if present
    if 'individual_question_analysis' in analysis_data and analysis_data['individual_question_analysis']:
        doc.add_heading('Detailed Question Analysis', level=1)

        for i, question in enumerate(analysis_data['individual_question_analysis'], 1):
            question_text = question.get('question_text', 'N/A')
            validity = question.get('validity', 'N/A')
            reason = question.get('reason', 'N/A')
            table_number = question.get('table_number', 'N/A')
            item_number = question.get('item_number', 'N/A')
            variable_name = question.get('variable_name', 'N/A')

            # Create a heading for each question using table and item numbers
            question_heading = doc.add_heading(f'Question Table {table_number} - {item_number}', level=2)
            if validity.lower() == 'not valid':
                from docx.shared import RGBColor
                question_heading.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red for invalid questions
            else:
                from docx.shared import RGBColor
                question_heading.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green for valid questions

            # Add question details in a table for better organization
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Attribute'
            hdr_cells[1].text = 'Value'

            # Add question details
            row_cells = table.add_row().cells
            row_cells[0].text = 'Table Number'
            row_cells[1].text = str(table_number)

            row_cells = table.add_row().cells
            row_cells[0].text = 'Item Number'
            row_cells[1].text = str(item_number)

            row_cells = table.add_row().cells
            row_cells[0].text = 'Variable Name'
            row_cells[1].text = variable_name

            row_cells = table.add_row().cells
            row_cells[0].text = 'Question Text'
            row_cells[1].text = question_text

            row_cells = table.add_row().cells
            row_cells[0].text = 'Validity'
            row_cells[1].text = validity
            if validity.lower() == 'not valid':
                from docx.shared import RGBColor
                row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red
                row_cells[1].paragraphs[0].runs[0].font.bold = True
            else:
                from docx.shared import RGBColor
                row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green
                row_cells[1].paragraphs[0].runs[0].font.bold = True

            row_cells = table.add_row().cells
            row_cells[0].text = 'Reason'
            row_cells[1].text = reason

            # Add alternative question if present and the question is not valid
            alternative_question = question.get('alternative_question', '')
            if alternative_question and validity.lower() == 'not valid':
                row_cells = table.add_row().cells
                row_cells[0].text = 'Suggested Alternative'
                row_cells[1].text = alternative_question

            # Add duplicate information if present
            duplicates_with = question.get('duplicates_with', [])
            if duplicates_with:
                row_cells = table.add_row().cells
                row_cells[0].text = 'Duplicates With'
                dup_text = ""
                for j, dup in enumerate(duplicates_with):
                    dup_table = dup.get('table_number', 'N/A')
                    dup_item = dup.get('item_number', 'N/A')
                    dup_text += f"Table {dup_table}, Item {dup_item}"
                    if j < len(duplicates_with) - 1:
                        dup_text += "; "
                row_cells[1].text = dup_text

            doc.add_paragraph("")  # Empty line for spacing

    # Add overall assessment
    doc.add_heading('Overall Assessment', level=1)
    if 'overall_assessment' in analysis_data:
        doc.add_paragraph(analysis_data['overall_assessment'])

    # Add general recommendations
    doc.add_heading('Recommendations', level=1)
    if 'recommendations' in analysis_data and analysis_data['recommendations']:
        for i, rec in enumerate(analysis_data['recommendations'], 1):
            p = doc.add_paragraph()
            p.add_run(f'{i}. ').bold = True
            p.add_run(rec)
    else:
        doc.add_paragraph('No specific recommendations provided.')

    # Add conclusion
    doc.add_heading('Conclusion', level=1)
    conclusion_para = doc.add_paragraph()
    conclusion_para.add_run('This report provides a comprehensive analysis of the survey questionnaire. ').bold = True
    conclusion_para.add_run('Based on the analysis, the survey has ')
    if valid_percentage >= 80:
        conclusion_para.add_run('good').bold = True
        conclusion_para.add_run(' quality with most questions being valid.')
    elif valid_percentage >= 60:
        conclusion_para.add_run('moderate').bold = True
        conclusion_para.add_run(' quality with some questions requiring attention.')
    else:
        conclusion_para.add_run('poor').bold = True
        conclusion_para.add_run(' quality with many questions needing revision.')

    # Save to temporary file
    temp_dir = tempfile.mkdtemp()
    temp_path = os.path.join(temp_dir, f"quality_report_{filename.replace('.txt', '').replace('.json', '').replace('.docx', '').replace('.pdf', '').replace('.csv', '')}.docx")
    doc.save(temp_path)

    return temp_path

if __name__ == "__main__":
    main()